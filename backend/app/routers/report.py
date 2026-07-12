from ..store import get_project_data
from ..dependencies import get_session_id, get_project_id
from fastapi import Depends, APIRouter, HTTPException
from ..store import get_project_data
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Dict, Any, List, Optional
import pandas as pd
import numpy as np
import os
import json
from .excel_report import generate_excel_report

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter()

class ReportConfig(BaseModel):
    format: str
    project_config: Dict[str, Any]
    auditLogs: Optional[List[Dict[str, Any]]] = []
    cachedResults: Optional[Dict[str, Any]] = {}

def add_audit_logs_to_word(doc, logs: List[Dict[str, Any]]):
    doc.add_heading("1. 데이터 전처리 및 분석 설정 기록 (Audit Trail)", level=1)
    if not logs:
        doc.add_paragraph("기록된 이력이 없습니다.")
        return
        
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '시간'
    hdr_cells[1].text = '단계'
    hdr_cells[2].text = '액션'
    hdr_cells[3].text = '상세 내용'
    
    for log in logs:
        row_cells = table.add_row().cells
        row_cells[0].text = str(log.get('timestamp', ''))
        row_cells[1].text = str(log.get('step', ''))
        row_cells[2].text = str(log.get('action', ''))
        
        details = log.get('details', '')
        if isinstance(details, dict) or isinstance(details, list):
            details_str = json.dumps(details, ensure_ascii=False)
        else:
            details_str = str(details)
            
        row_cells[3].text = details_str

def add_analysis_results_to_word(doc, cached_results: Dict[str, Any]):
    doc.add_heading("2. 분석 결과 및 해석", level=1)
    if not cached_results:
        doc.add_paragraph("저장된 분석 결과가 없습니다.")
        return
        
    for key, cache in cached_results.items():
        doc.add_heading(f"분석: {key.upper()}", level=2)
        
        settings = cache.get("settings", {})
        if settings:
            doc.add_heading("사용된 설정(Options)", level=3)
            for sk, sv in settings.items():
                doc.add_paragraph(f"- {sk}: {sv}")
        
        interpretation = cache.get("interpretation", "")
        if interpretation:
            doc.add_heading("결과 해석(Interpretation)", level=3)
            doc.add_paragraph(interpretation)
            
        # raw data is too complex to neatly table in python without specific schemas,
        # so we will rely on interpretation for Word, and put raw results in Excel.

@router.post("/full-report")
async def generate_full_report(request: ReportConfig, session_id: str = Depends(get_session_id), project_id: str = Depends(get_project_id)):
    df = get_project_data(project_id, session_id)
    original_df = get_project_data(project_id, session_id, original=True)
    if df is None:
        raise HTTPException(status_code=400, detail="데이터가 메모리에 존재하지 않습니다. 먼저 업로드해주세요.")
    
    if request.format == 'word':
        doc = Document()
        
        title = doc.add_heading("통계 분석 전체 보고서", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. Audit Logs
        add_audit_logs_to_word(doc, request.auditLogs)
        
        # 2. Cached Results (Interpretations & Settings)
        add_analysis_results_to_word(doc, request.cachedResults)
        
        # 파일 저장
        output_path = "data/Full_Report.docx"
        doc.save(output_path)
        return FileResponse(path=output_path, filename="Full_Report.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
    elif request.format == 'excel':
        output_path = "data/Full_Report.xlsx"
        try:
            generate_excel_report(output_path, request, df, original_df)
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=str(e))
            
        return FileResponse(path=output_path, filename="Full_Report.xlsx", media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    else:
        raise HTTPException(status_code=400, detail="지원하지 않는 포맷입니다.")
